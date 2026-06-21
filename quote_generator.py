#!/usr/bin/env python3
"""
TeamLogic IT of Sherman-Celina — Quote Generator
=================================================

Builds branded Word (.docx) quotes in two modes:

  1) MRR / Managed Services  -> reads the "Managed Services Estimator v4.4"
     workbook and pulls the proposal-safe (rounded) figures from the Results
     sheet, then writes a managed-services proposal.

  2) Time & Materials        -> takes an hourly rate (+ after-hours multiplier,
     minimum, and increment) and writes a T&M proposal.

Runs as an interactive Q&A in the terminal:

    python3 quote_generator.py

Dependencies: python-docx, openpyxl
    pip install python-docx openpyxl --break-system-packages

Optional (for recalculating an estimator that wasn't last saved by Excel,
and for PDF export): LibreOffice (`soffice` on PATH).
"""

import os
import sys
import shutil
import subprocess
from datetime import date, datetime

try:
    from openpyxl import load_workbook
except ImportError:
    sys.exit("Missing dependency: openpyxl.  Run: pip install openpyxl --break-system-packages")

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ----------------------------------------------------------------------------
# Brand constants  (TeamLogic IT brand standards)
# ----------------------------------------------------------------------------
GREEN = "008851"   # PMS 348 — primary
BLUE  = "0095DA"   # PMS 2925 — accent / rules only
GRAY  = "636467"   # Cool Gray 10
DARK  = "222222"
LIGHT = "EAF4EE"   # light green fill for table bands
FONT  = "Montserrat"

HERE = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(HERE, "logo.png")

CONTACT = "Brian Camp  •  972-807-0878  •  bcamp@teamlogicit.com  •  teamlogicit.com/Sherman-CelinaTX"
COPYRIGHT = "©2026 TeamLogic, Inc. All rights reserved. Business Confidential."


# ----------------------------------------------------------------------------
# Low-level docx helpers
# ----------------------------------------------------------------------------
def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        m.append(node)
    tcPr.append(m)


def _bottom_border(paragraph, color, size=6):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _left_bar(paragraph, color, size=18):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    pbdr.append(left)
    pPr.append(pbdr)


def run(paragraph, text, size=10.5, color=DARK, bold=False, italic=False):
    r = paragraph.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    r.italic = italic
    return r


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(7)
    run(p, text, size=13, color=GREEN, bold=True)
    _bottom_border(p, GREEN, size=6)
    return p


def body(doc, runs, after=6, align=None, indent=None):
    """runs: list of (text, kwargs) tuples or a single string."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if isinstance(runs, str):
        run(p, runs)
    else:
        for text, kw in runs:
            run(p, text, **kw)
    return p


def bullet(doc, runs, after=4):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(after)
    if isinstance(runs, str):
        run(p, runs)
    else:
        for text, kw in runs:
            run(p, text, **kw)
    return p


def letterhead(doc, subtitle):
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.add_run().add_picture(LOGO_PATH, width=Inches(2.2))
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    run(p, "TeamLogic IT of Sherman-Celina", size=16, color=GREEN, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    run(p, subtitle, size=12, color=DARK)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(9)
    run(p, CONTACT, size=9, color=GRAY)
    _bottom_border(p, BLUE, size=12)


def footer_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run(p, COPYRIGHT, size=7.5, color=GRAY, italic=True)


def styled_table(doc, headers, col_widths_in):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    for j, (h, w) in enumerate(zip(headers, col_widths_in)):
        cell = table.rows[0].cells[j]
        cell.width = Inches(w)
        _set_cell_bg(cell, GREEN)
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j else WD_ALIGN_PARAGRAPH.LEFT
        run(para, h, size=10, color="FFFFFF", bold=True)
    return table


def add_row(table, cells, col_widths_in, fill=None):
    """cells: list of (text, {kwargs, 'align':..}) ; widths align to columns."""
    row = table.add_row()
    for j, (spec, w) in enumerate(zip(cells, col_widths_in)):
        cell = row.cells[j]
        cell.width = Inches(w)
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if fill:
            _set_cell_bg(cell, fill)
        text, kw = spec
        align = kw.pop("align", None)
        para = cell.paragraphs[0]
        if align:
            para.alignment = align
        run(para, text, **kw)
    return row


def section_band(table, label, ncols):
    """Full-width light-green band row spanning all columns."""
    row = table.add_row()
    a = row.cells[0]
    b = row.cells[-1]
    a.merge(b)
    cell = row.cells[0]
    _set_cell_bg(cell, LIGHT)
    _set_cell_margins(cell, top=50, bottom=50)
    run(cell.paragraphs[0], label, size=10, color=GREEN, bold=True)


# ----------------------------------------------------------------------------
# Estimator reader
# ----------------------------------------------------------------------------
class EstimatorData:
    """Pulls proposal-safe values from the Managed Services Estimator Results sheet."""

    def __init__(self, prospect, environment, onsite_remote_mrr,
                 remote_mrr, onboarding, line_items):
        self.prospect = prospect
        self.environment = environment          # dict label -> value
        self.onsite_remote_mrr = onsite_remote_mrr
        self.remote_mrr = remote_mrr
        self.onboarding = onboarding
        self.line_items = line_items            # list of (category, included_bool)


def _find_value_right_of(ws, label, max_row=60, max_col=20):
    """Find a cell whose text == label, return the first non-empty value to its right."""
    for r in ws.iter_rows(min_row=1, max_row=max_row, max_col=max_col):
        for c in r:
            if isinstance(c.value, str) and c.value.strip() == label:
                for cc in ws[c.row][c.column:]:
                    if cc.value not in (None, ""):
                        return cc.value
    return None


def _find_money_after_label(ws, label, max_row=60, max_col=20, exclude=None):
    """Find row(s) whose text contains label (but not `exclude`), return the
    numeric values on that row. Prefers positive values (a cached 0 usually
    means the file wasn't saved by Excel)."""
    found = []
    for r in ws.iter_rows(min_row=1, max_row=max_row, max_col=max_col):
        texts = [c.value for c in r if isinstance(c.value, str)]
        if any(label in t for t in texts) and not (
                exclude and any(exclude in t for t in texts)):
            nums = [c.value for c in r if isinstance(c.value, (int, float))]
            if nums:
                found.extend(nums)
    positives = [n for n in found if n > 0]
    return positives if positives else found


def recalc_if_possible(path):
    """Deliberately a no-op on the estimator.

    The estimator's Results sheet uses formulas whose values are cached by
    Excel on save. Re-running a headless recalc (LibreOffice / recalc.py) can
    zero-out those cached values because the workbook depends on Excel-specific
    evaluation. So we trust the values Excel last saved and do NOT recalc here.

    If you ever need a fresh recalc, open and save the file in Excel first.
    """
    return path


def read_estimator(path):
    path = recalc_if_possible(path)
    wb = load_workbook(path, data_only=True)
    if "Results" not in wb.sheetnames:
        raise ValueError("This workbook has no 'Results' sheet — is it the Managed Services Estimator?")
    ws = wb["Results"]

    prospect = _find_value_right_of(ws, "Prospect:") or "Prospect"

    # Environment inputs live in column B (label) / C (value)
    env_labels = ["Servers", "Computers-All", "Mobile Devices", "Employees",
                  "Firewalls", "Domains", "Mailboxes", "Incidents", "TB"]
    environment = {}
    for r in ws.iter_rows(min_row=1, max_row=40, max_col=4):
        for c in r:
            if isinstance(c.value, str) and c.value.strip() in env_labels:
                val = ws.cell(row=c.row, column=c.column + 1).value
                if val not in (None, ""):
                    environment[c.value.strip()] = val

    # Proposal-safe rounded figures
    onsite = _find_money_after_label(ws, "Managed Services with Onsite & Remote Support")
    remote = _find_money_after_label(ws, "Managed Services with Remote Support",
                                     exclude="Onsite")
    onboard = _find_money_after_label(ws, "Onboarding:")

    onsite_remote_mrr = max(onsite) if onsite else None
    remote_mrr = max(remote) if remote else None
    onboarding = max(onboard) if onboard else None

    if onsite_remote_mrr in (None, 0) and remote_mrr in (None, 0):
        print("\n  ⚠  Warning: the estimator's MRR cells read as empty/zero.")
        print("     This usually means the workbook wasn't last saved by Excel.")
        print("     Open the estimator in Excel, save it, and try again.\n")

    # Which service categories are included (cost > 0 in the monthly recurring block)
    cat_rows = {
        "Managed Services Essentials": None,
        "Backup": None, "Help Desk": None,
        "Cybersecurity Essentials": None,
        "Onsite Support": None,
    }
    line_items = []
    for r in ws.iter_rows(min_row=1, max_row=40, max_col=12):
        for c in r:
            if isinstance(c.value, str) and c.value.strip() in cat_rows:
                # the numeric monthly cost is in column I (9) on the same row
                cost = ws.cell(row=c.row, column=9).value
                included = isinstance(cost, (int, float)) and cost > 0
                line_items.append((c.value.strip(), included))

    return EstimatorData(prospect, environment, onsite_remote_mrr,
                         remote_mrr, onboarding, line_items)


# ----------------------------------------------------------------------------
# Document builders
# ----------------------------------------------------------------------------
MS_SERVICES = [
    ("Managed Services Essentials", [
        ("Remote Monitoring & Management", "24/7 monitoring and management of all workstations (RMM agent on every device)."),
        ("Patch Management", "Operating-system and third-party software patching, fully managed."),
        ("Network Operations Center", "Proactive NOC oversight of computers and infrastructure."),
        ("Cloud Backup", "Cloud-based data protection covering workstations and business data."),
        ("Help Desk Support", "U.S.-based help desk for day-to-day user support and issue resolution."),
    ]),
    ("Cybersecurity Essentials", [
        ("Endpoint & Identity Protection", "Next-gen antivirus + EDR + 24/7 Managed Detection & Response."),
        ("DNS Filtering", "Web/DNS-layer filtering to block malicious and inappropriate sites."),
        ("Password Manager", "Business password manager for all staff members."),
        ("Email Encryption & Threat Protection", "Encrypted email with threat protection and archiving."),
        ("Security Awareness Training", "Ongoing phishing simulation and security awareness training."),
        ("Dark Web Monitoring", "Continuous monitoring of the domain for compromised credentials."),
    ]),
]


def save_doc(doc, outfile):
    """Save the document, falling back to a timestamped name if the target is locked.

    The usual cause of a locked target is that a previously generated copy is still
    open in Word, which holds an exclusive lock and makes doc.save() raise
    PermissionError. Rather than crash, we save alongside it and tell the user.
    """
    try:
        doc.save(outfile)
        return outfile
    except PermissionError:
        base, ext = os.path.splitext(outfile)
        alt = f"{base}_{datetime.now():%H%M%S}{ext}"
        name = os.path.basename(outfile)
        print(f"\n  ⚠  Could not overwrite '{name}' — it looks like it's open in Word.")
        print(f"     Saving as '{os.path.basename(alt)}' instead. Close the open copy to")
        print("     overwrite it directly next time.")
        doc.save(alt)
        return alt


def base_document():
    doc = Document()
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(10.5)
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    return doc


def build_managed_quote(data, price_mode, outfile):
    """price_mode: 'both' | 'onsite' | 'remote'"""
    doc = base_document()
    letterhead(doc, "Managed Services Quote")

    body(doc, [("Prepared for:  ", {"bold": True}), (data.prospect, {})], after=1)
    if data.environment:
        env = "  •  ".join(f"{v} {k.lower()}" for k, v in data.environment.items()
                           if isinstance(v, (int, float)) and v)
        if env:
            body(doc, [("Environment:  ", {"bold": True}), (env, {})], after=1)
    body(doc, [("Date:  ", {"bold": True}),
               (date.today().strftime("%B %d, %Y"), {})], after=8)

    body(doc, f"Thank you for the opportunity to support {data.prospect}. The plan below "
              "wraps your computers, email, and data in proactive management and layered "
              "security — so your team stays focused on the work while we keep the technology "
              "secure, monitored, and running.")

    # Services table
    heading(doc, "Services Included")
    widths = [2.4, 4.1]  # sum 6.5" (usable page width)
    t = styled_table(doc, ["Service", "What's Included"], widths)
    for category, items in MS_SERVICES:
        section_band(t, category, 2)
        for name, detail in items:
            add_row(t, [(name, {"size": 10, "bold": True}),
                        (detail, {"size": 10, "color": GRAY})], widths)

    # Pricing
    heading(doc, "Monthly Investment")
    body(doc, [("Both options include the full stack of Managed Services and Cybersecurity "
                "Essentials above — the difference is whether on-site support is included "
                "alongside remote support.", {"size": 10, "color": GRAY})], after=7) \
        if price_mode == "both" else None

    pw = [3.5, 1.0, 2.0]  # sum 6.5" (usable page width)
    pt = styled_table(doc, ["Plan", "Support", "Monthly Price"], pw)
    show_onsite = price_mode in ("both", "onsite")
    show_remote = price_mode in ("both", "remote")
    if show_onsite and data.onsite_remote_mrr is not None:
        add_row(pt, [("Managed Services — Onsite & Remote Support",
                      {"size": 11, "bold": True, "color": GREEN}),
                     ("On-site + remote", {"size": 9, "color": GRAY,
                      "align": WD_ALIGN_PARAGRAPH.CENTER}),
                     (f"${data.onsite_remote_mrr:,.0f} / mo",
                      {"size": 13, "bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER})],
                pw, fill=LIGHT)
    if show_remote and data.remote_mrr is not None:
        add_row(pt, [("Managed Services — Remote Support",
                      {"size": 11, "bold": True, "color": GREEN}),
                     ("Remote", {"size": 9, "color": GRAY,
                      "align": WD_ALIGN_PARAGRAPH.CENTER}),
                     (f"${data.remote_mrr:,.0f} / mo",
                      {"size": 13, "bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER})],
                pw)

    # Onboarding
    if data.onboarding:
        heading(doc, "One-Time Onboarding")
        body(doc, [("A one-time onboarding fee of ", {}),
                   (f"${data.onboarding:,.0f}", {"bold": True}),
                   (" covers initial setup, agent deployment, security configuration, "
                    "and documentation of your environment.", {})])

    # Next steps
    heading(doc, "Next Steps")
    body(doc, "To move forward, two documents will be put in place: a Master Services "
              "Agreement (MSA) that sets the overall terms of our relationship, and a "
              "Statement of Work (SOW) that details the services and pricing in this quote. "
              "Once both are signed, we'll schedule onboarding at a time that works for you.")
    body(doc, "To get started or walk through any line item, call Brian Camp at "
              "972-807-0878 or email bcamp@teamlogicit.com.")
    footer_line(doc)
    return save_doc(doc, outfile)


def build_tm_quote(rate, after_mult, minimum_hrs, increment_min,
                   prospect, bh_window, outfile):
    after_rate = rate * after_mult
    doc = base_document()
    letterhead(doc, "Time & Materials Service Quote")

    body(doc, [("Prepared for:  ", {"bold": True}), (prospect, {})], after=1)
    body(doc, [("Date:  ", {"bold": True}),
               (date.today().strftime("%B %d, %Y"), {})], after=8)

    body(doc, "This quote covers IT support on a time-and-materials basis — a flexible, "
              "pay-as-you-go option with no monthly commitment. You're billed only for the "
              "time we spend resolving your issues and completing your projects, at the rates below.")

    heading(doc, "Hourly Rates")
    rw = [2.4, 2.6, 1.5]  # must sum to <= 6.5" (usable page width) so all columns render in Word
    rt = styled_table(doc, ["Service Window", "When It Applies", "Hourly Rate"], rw)
    add_row(rt, [("Business Hours", {"size": 11, "bold": True, "color": GREEN}),
                 (bh_window, {"size": 10, "color": GRAY}),
                 (f"${rate:,.0f} / hr", {"size": 12, "bold": True,
                  "align": WD_ALIGN_PARAGRAPH.CENTER})], rw, fill=LIGHT)
    pct = int(round((after_mult - 1) * 100))
    add_row(rt, [("After Hours", {"size": 11, "bold": True, "color": GREEN}),
                 (f"Evenings, weekends & holidays (+{pct}%)", {"size": 10, "color": GRAY}),
                 (f"${after_rate:,.0f} / hr", {"size": 12, "bold": True,
                  "align": WD_ALIGN_PARAGRAPH.CENTER})], rw)

    heading(doc, "How Time Is Billed")
    inc_label = f"{increment_min}-minute"
    bullet(doc, [(f"{minimum_hrs:g}-hour minimum", {"bold": True}),
                 (" on every support request.", {})])
    bullet(doc, [(f"Billed in {inc_label} increments", {"bold": True}),
                 (" after the first hour.", {})])
    bullet(doc, [("After-hours rate", {"bold": True}),
                 (f" applies outside {bh_window.lower()}, and on holidays, "
                  f"at {after_mult:g}× the business-hours rate.", {})])
    bullet(doc, [("Materials, hardware, and third-party software", {"bold": True}),
                 (" (if any) are billed at cost and itemized separately.", {})])

    # Billing examples (computed from the actual rules)
    heading(doc, "Billing Examples")
    ew = [2.3, 1.6, 1.3, 1.3]  # sum 6.5" so the Total column stays on the page in Word
    et = styled_table(doc, ["Example", "Time Worked", "Billed As", "Total"], ew)

    def round_increment(hours):
        """Apply minimum then round up to the increment after the first hour."""
        step = increment_min / 60.0
        if hours <= minimum_hrs:
            return minimum_hrs
        over = hours - minimum_hrs
        import math
        steps = math.ceil(round(over / step, 6))
        return minimum_hrs + steps * step

    ex = [
        ("Quick fix, business hours", 40/60, rate),
        ("Project work, business hours", 1 + 50/60, rate),
        ("After-hours emergency", 1 + 15/60, after_rate),
    ]
    for label, worked_hrs, r in ex:
        billed = round_increment(worked_hrs)
        total = billed * r
        wh = int(worked_hrs)
        wm = int(round((worked_hrs - wh) * 60))
        worked_str = (f"{wh} hr " if wh else "") + (f"{wm} min" if wm else "").strip()
        worked_str = worked_str.strip() or f"{int(worked_hrs*60)} min"
        billed_str = f"{billed:g} hr" + ("s" if billed != 1 else "")
        if billed == minimum_hrs and worked_hrs <= minimum_hrs:
            billed_str += " (min.)"
        add_row(et, [(label, {"size": 10}),
                     (worked_str, {"size": 10, "color": GRAY}),
                     (billed_str, {"size": 10, "color": GRAY}),
                     (f"${total:,.2f}", {"size": 11, "bold": True,
                      "align": WD_ALIGN_PARAGRAPH.CENTER})], ew)

    heading(doc, "Next Steps")
    body(doc, "To put this in place, two documents are required: a signed Master Services "
              "Agreement (MSA) setting the overall terms, and a Statement of Work (SOW) "
              "covering these time-and-materials rates and billing terms. Once both are "
              "signed, you can request support whenever you need it and we'll schedule promptly.")
    body(doc, "To get started or ask any questions, contact Brian Camp at 972-807-0878 or "
              "bcamp@teamlogicit.com. If your support needs grow, we're glad to discuss a "
              "flat-rate managed services plan that can lower your overall cost.")
    footer_line(doc)
    return save_doc(doc, outfile)


# ----------------------------------------------------------------------------
# PDF export (optional)
# ----------------------------------------------------------------------------
def open_file(path):
    """Open a file in the OS default application (Word for .docx on Windows)."""
    try:
        if sys.platform.startswith("win"):
            os.startfile(path)  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.run(["open", path], check=False)
        else:
            subprocess.run(["xdg-open", path], check=False)
    except Exception as e:
        print(f"  (Could not auto-open the file: {e})")


def export_pdf(docx_path):
    if shutil.which("soffice") is None:
        print("  (LibreOffice not found — skipping PDF export.)")
        return None
    outdir = os.path.dirname(os.path.abspath(docx_path))
    try:
        subprocess.run(["soffice", "--headless", "--convert-to", "pdf",
                        "--outdir", outdir, docx_path],
                       check=True, capture_output=True, timeout=120)
        pdf = os.path.splitext(docx_path)[0] + ".pdf"
        return pdf if os.path.exists(pdf) else None
    except Exception as e:
        print(f"  (PDF export failed: {e})")
        return None


# ----------------------------------------------------------------------------
# Interactive prompts
# ----------------------------------------------------------------------------
def ask(prompt, default=None):
    suffix = f" [{default}]" if default not in (None, "") else ""
    val = input(f"{prompt}{suffix}: ").strip()
    return val if val else (default if default is not None else "")


def ask_float(prompt, default):
    while True:
        raw = ask(prompt, default)
        try:
            return float(raw)
        except ValueError:
            print("  Please enter a number.")


def ask_choice(prompt, options):
    print(prompt)
    for i, (key, label) in enumerate(options, 1):
        print(f"  {i}) {label}")
    while True:
        raw = input("Choose: ").strip()
        if raw.isdigit() and 1 <= int(raw) <= len(options):
            return options[int(raw) - 1][0]
        print("  Invalid choice.")


def slugify(name):
    keep = "".join(c if c.isalnum() else "_" for c in name)
    while "__" in keep:
        keep = keep.replace("__", "_")
    return keep.strip("_") or "Quote"


def ensure_outdir(outdir):
    """Normalize an output-folder entry and make sure it exists.

    Tolerates a quoted path and the common slip of pasting a *file* path (e.g. the
    estimator workbook) where a folder is expected — in that case we use the file's
    containing folder instead of crashing in os.makedirs.
    """
    outdir = os.path.expanduser(outdir.strip().strip('"').strip("'"))
    if os.path.isfile(outdir):
        folder = os.path.dirname(outdir)
        print(f"  (That's a file, not a folder — using its folder instead: {folder})")
        outdir = folder
    outdir = outdir or os.getcwd()
    os.makedirs(outdir, exist_ok=True)
    return outdir


def main():
    # Windows consoles default to cp1252, which can't encode the ✓/⚠/• glyphs we print.
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass
    print("=" * 60)
    print("  TeamLogic IT of Sherman-Celina — Quote Generator")
    print("=" * 60)

    mode = ask_choice("\nWhat kind of quote?",
                      [("mrr", "Managed Services / MRR (from the estimator workbook)"),
                       ("tm", "Time & Materials (hourly rate)")])

    if mode == "mrr":
        path = ask("\nPath to the Managed Services Estimator (.xlsm)")
        # Windows "Copy as path" wraps the path in double quotes; strip surrounding quotes.
        path = os.path.expanduser(path.strip().strip('"').strip("'"))
        if not os.path.exists(path):
            sys.exit(f"File not found: {path}")
        print("\nReading estimator…")
        try:
            data = read_estimator(path)
        except Exception as e:
            sys.exit(f"Could not read estimator: {e}")

        print(f"  Prospect: {data.prospect}")
        if data.onsite_remote_mrr:
            print(f"  Onsite & Remote MRR: ${data.onsite_remote_mrr:,.0f}/mo")
        if data.remote_mrr:
            print(f"  Remote MRR:          ${data.remote_mrr:,.0f}/mo")
        if data.onboarding:
            print(f"  Onboarding:          ${data.onboarding:,.0f}")

        # let the user override the auto-detected prospect name
        override = ask("\nProspect name on the quote", data.prospect)
        data.prospect = override or data.prospect

        price_mode = ask_choice("\nWhich price(s) to show?",
                                [("both", "Both options side by side"),
                                 ("onsite", "Onsite & Remote only"),
                                 ("remote", "Remote only")])

        default_outdir = os.path.dirname(os.path.abspath(path))
        quote_name = f"Quote_{slugify(data.prospect)}.docx"

        def builder(target):
            return build_managed_quote(data, price_mode, target)

    else:  # time & materials
        prospect = ask("\nProspect / client name", "Prospective Client")
        rate = ask_float("Business-hours rate ($/hr)", "150")
        mult_pct = ask_float("After-hours uplift (%)", "50")
        after_mult = 1 + mult_pct / 100.0
        minimum_hrs = ask_float("Minimum billable hours", "1")
        increment_min = ask_float("Billing increment after first hour (minutes)", "30")
        bh_window = ask("Business-hours window",
                        "Monday–Friday, 8:00 AM – 5:00 PM")

        default_outdir = os.getcwd()
        quote_name = f"Quote_TandM_{slugify(prospect)}.docx"

        def builder(target):
            return build_tm_quote(rate, after_mult, minimum_hrs, increment_min,
                                  prospect, bh_window, target)

    outdir = ensure_outdir(ask("\nOutput folder", default_outdir))
    want_pdf = ask("Also export a PDF? (y/n)", "n").lower().startswith("y")
    outfile = builder(os.path.join(outdir, quote_name))

    print(f"\n✓ Created: {outfile}")
    if want_pdf:
        pdf = export_pdf(outfile)
        if pdf:
            print(f"✓ Created: {pdf}")
    print("Done.")
    open_file(outfile)


if __name__ == "__main__":
    main()
